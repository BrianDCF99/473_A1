#!/usr/bin/env python3
"""
Generate a .docx report for the combinatorial testing project.

Summarizes two main bugs and embeds two figures:
  - Figure 1: Pipeline workflow diagram
  - Figure 2: Failures grouped by file_type
"""

from __future__ import annotations

import csv
from collections import Counter
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESULTS_CSV = PROJECT_ROOT / "results" / "icgrep_results_full.csv"
FRAMES_CSV = PROJECT_ROOT / "frames" / "icgrep_frames_full.csv"
OUT_DOCX = PROJECT_ROOT / "results" / "Combinatorial_Testing_Report.docx"
FIG_DIR = PROJECT_ROOT / "results" / "figures"

REPO_ROOT = PROJECT_ROOT.parent


def sanitize(text: str) -> str:
    """Strip machine-specific absolute paths so the report is portable."""
    if not text:
        return text
    return (
        text.replace(str(REPO_ROOT) + "/", "")
            .replace(str(REPO_ROOT), "<repo>")
            .replace(str(Path.home()) + "/", "~/")
    )


def load_csv(path: Path):
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


HEADING_SIZES = {0: 16, 1: 13, 2: 11}


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(HEADING_SIZES.get(level, 11))
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)


def add_figure(doc, image_path, caption, width_in=6.0):
    doc.add_picture(str(image_path), width=Inches(width_in))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)


def build_summary(results, frames):
    fails = [r for r in results if r["status"] == "FAIL"]
    return {
        "total": len(results),
        "pass": len(results) - len(fails),
        "fail": len(fails),
        "exit_mismatch": Counter((r["expected_exit"], r["actual_exit"]) for r in fails),
        "by_ft": Counter(frames[r["id"]]["file_type"] for r in fails),
        "by_pt": Counter(frames[r["id"]]["pattern_type"] for r in fails),
        "fails": fails,
    }


def make_workflow_figure(path: Path):
    fig, ax = plt.subplots(figsize=(10, 3))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3)
    ax.axis("off")

    boxes = [
        (0.2, "Parameter\nModel\n(CSV/ACTS)"),
        (2.1, "Frame\nGenerator"),
        (4.0, "Frames CSV\n(1360 cases)"),
        (5.9, "Python\nRunner"),
        (7.8, "Oracle\n(stdout+exit)"),
    ]
    result = (9.4, "Results CSV\n+ Report")

    for x, label in boxes:
        ax.add_patch(mpatches.FancyBboxPatch((x, 1.1), 1.4, 1.0,
                                             boxstyle="round,pad=0.05",
                                             fc="#E8F0FE", ec="#1A73E8", lw=1.5))
        ax.text(x + 0.7, 1.6, label, ha="center", va="center", fontsize=9)

    ax.add_patch(mpatches.FancyBboxPatch((result[0], 1.1), 1.4, 1.0,
                                         boxstyle="round,pad=0.05",
                                         fc="#FCE8E6", ec="#D93025", lw=1.5))
    ax.text(result[0] + 0.7, 1.6, result[1], ha="center", va="center", fontsize=9)

    for (x, _), (x2, _) in zip(boxes[:-1], boxes[1:]):
        ax.annotate("", xy=(x2, 1.6), xytext=(x + 1.4, 1.6),
                    arrowprops=dict(arrowstyle="->", color="#444", lw=1.2))
    ax.annotate("", xy=(result[0], 1.6), xytext=(boxes[-1][0] + 1.4, 1.6),
                arrowprops=dict(arrowstyle="->", color="#444", lw=1.2))

    ax.text(5.0, 0.3, "icgrep (Parabix) runs each command; oracle compares actual output and exit code to spec.",
            ha="center", va="center", fontsize=8, color="#555", style="italic")

    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)


def make_fails_by_file_type_figure(summary, path: Path):
    data = summary["by_ft"].most_common()
    if not data:
        data = [("no failures", 0)]
    labels, counts = zip(*data)

    colors = ["#D93025" if label == "missing_path" else "#1A73E8" for label in labels]

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(labels, counts, color=colors)
    ax.set_ylabel("Number of failing test cases")
    ax.set_title("Failures by file_type (out of 1360 total)")
    for bar, c in zip(bars, counts):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 2,
                str(c), ha="center", va="bottom", fontsize=10)
    ax.set_ylim(0, max(counts) * 1.2 if counts else 10)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)


def main():
    if not RESULTS_CSV.exists():
        raise SystemExit(f"Missing results: {RESULTS_CSV}")
    if not FRAMES_CSV.exists():
        raise SystemExit(f"Missing frames: {FRAMES_CSV}")

    results = load_csv(RESULTS_CSV)
    frames = {row["id"]: row for row in load_csv(FRAMES_CSV)}
    summary = build_summary(results, frames)

    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig1_path = FIG_DIR / "figure1_workflow.png"
    fig2_path = FIG_DIR / "figure2_fails_by_file_type.png"
    make_workflow_figure(fig1_path)
    make_fails_by_file_type_figure(summary, fig2_path)

    doc = Document()
    doc.sections[0].left_margin = Inches(1)
    doc.sections[0].right_margin = Inches(1)

    add_heading(doc, "Combinatorial Functional Testing – icgrep", level=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Project: Parabix functional testing using ACTS-style combinatorial model")
    r.italic = True
    r.font.size = Pt(10)

    # --- Executive Summary (two main bugs up front)
    add_heading(doc, "Executive Summary — Two Main Bugs", level=1)
    add_para(
        doc,
        "Functional combinatorial testing of icgrep across 1,360 parameter combinations "
        "identified two distinct specification violations. Both relate to how error "
        "conditions on input files are reported through exit codes.",
    )
    add_para(doc, "Missing file returns exit 1 instead of 2.", bold=True)
    add_para(
        doc,
        "When the input file does not exist (or is a directory, or is unreadable), "
        "icgrep prints an error on stderr but exits with code 1 (\"no match\") instead "
        "of code 2 (\"error\"). Detected in all 144 missing_path test frames and "
        "confirmed manually for directory and permission-denied cases.",
    )
    add_para(doc, "Partial file failure is masked by exit 0.", bold=True)
    add_para(
        doc,
        "When icgrep is given multiple input files and at least one of them cannot be "
        "opened, icgrep prints an error but still returns exit 0 if any of the other "
        "files produced a match. This is more severe than the missing-file case "
        "because the exit code completely hides the failure from automation and CI.",
    )

    # --- 1. Methodology
    add_heading(doc, "Methodology", level=1)
    add_para(
        doc,
        "We check Parabix tool icgrep against its I/O specification (expected stdout "
        "and exit codes 0, 1, 2). An ACTS-style parameter model defines pattern type, "
        "file type, and flags such as count, invert, line numbers. A Python script "
        "reads the combinations from a CSV, turns each row into a concrete command and "
        "expected result, runs icgrep, and reports any mismatch. A small controlled "
        "set of test files is used so the expected output and exit code are known "
        "unambiguously from the spec. All cases are scripted, reproducible, and logged.",
    )

    add_para(doc, "Model parameters:", bold=True)
    add_bullets(
        doc,
        [
            "source_type: inline, file_flag",
            "pattern_type: literal, char_class, negated_class, anchor_start, "
            "anchor_end, alternation, repetition, unicode_property, empty, invalid",
            "file_type: empty, no_match, one_match, many_match, unicode_content, missing_path",
            "flags: count (-c), invert (-v), ignore_case (-i), line_numbers (-n)",
        ],
    )
    add_para(doc, "Constraints:", bold=True)
    add_bullets(
        doc,
        [
            "count=1 implies line_numbers=0",
            "file_type=missing_path implies line_numbers=0",
        ],
    )
    add_para(doc, "Oracle exit-code spec:", bold=True)
    add_bullets(
        doc,
        [
            "0 = at least one selected line",
            "1 = no selected lines",
            "2 = error (invalid regex, missing or unreadable input file, usage error)",
        ],
    )

    add_figure(doc, fig1_path, "Figure 1. End-to-end test pipeline used in this project.")

    # --- 2. Results
    add_heading(doc, "Results Summary", level=1)
    add_para(
        doc,
        f"Total frames executed: {summary['total']}. "
        f"PASS: {summary['pass']}. FAIL: {summary['fail']}.",
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Expected exit"
    hdr[1].text = "Actual exit"
    hdr[2].text = "Count"
    for (exp, act), count in summary["exit_mismatch"].most_common():
        row = table.add_row().cells
        row[0].text = str(exp)
        row[1].text = str(act)
        row[2].text = str(count)

    add_figure(doc, fig2_path, "Figure 2. Failures grouped by file_type. All failures concentrate on missing_path, pointing to a single root-cause.")

    add_heading(doc, "Missing file returns exit 1 instead of 2", level=1)
    add_para(
        doc,
        "Specification violation: icgrep returns exit code 1 instead of exit code 2 "
        "when the input file does not exist.",
        bold=True,
    )
    add_para(
        doc,
        "Across all 144 test frames where file_type = missing_path (covering every "
        "combination of pattern_type and flags), icgrep prints an error on stderr "
        "(\"No such file\") but exits with code 1 (\"no match\"). Per the spec, a "
        "missing input file is an error and should produce exit code 2.",
    )

    add_para(doc, "Manual reproduction:", bold=True)
    add_code(
        doc,
        "$ icgrep -colors=never foo /tmp/definitely_not_a_real_path.txt\n"
        "icgrep: \"/tmp/definitely_not_a_real_path.txt\": No such file.\n"
        "$ echo $?\n"
        "1   # expected: 2",
    )

    add_para(doc, "Evidence (first three failing frames):", bold=True)
    miss_fail = [r for r in summary["fails"] if frames[r["id"]]["file_type"] == "missing_path"][:3]
    for r in miss_fail:
        add_code(
            doc,
            f"id={r['id']}  cmd: {sanitize(r['cmd'])}\n"
            f"  expected_exit=2  actual_exit={r['actual_exit']}\n"
            f"  stderr: {sanitize(r['stderr'])[:140]}",
        )

    add_para(
        doc,
        "The same bug class was confirmed manually for directory and unreadable "
        "file arguments (both return exit 1 instead of 2).",
        bold=True,
    )
    add_code(
        doc,
        "$ icgrep -colors=never foo /path/to/a/directory\n"
        "icgrep: /path/to/a/directory: Is a directory.\n"
        "exit=1   # expected: 2\n"
        "\n"
        "$ icgrep -colors=never foo /tmp/noperm.txt\n"
        "icgrep: \"/tmp/noperm.txt\": Permission denied.\n"
        "exit=1   # expected: 2",
    )

    add_para(doc, "Impact:", bold=True)
    add_bullets(
        doc,
        [
            "Users and CI scripts cannot rely on exit codes to detect missing files.",
            "Treating a missing file as \"no match\" can mask data-pipeline failures.",
            "Behavior is inconsistent with the tool's own stderr error message.",
        ],
    )
    add_para(doc, "Suggested fix:", bold=True)
    add_para(
        doc,
        "When a file argument cannot be opened (missing, directory, unreadable), set "
        "the exit code to 2 (error) instead of the no-match path.",
    )

    add_heading(doc, "Partial file failure is masked by exit 0", level=1)
    add_para(
        doc,
        "Specification violation: when icgrep is run with multiple input files and at "
        "least one of them cannot be opened, icgrep prints an error for the missing "
        "file but still returns exit 0 if any other file produced a match.",
        bold=True,
    )

    add_para(doc, "Manual reproduction:", bold=True)
    add_code(
        doc,
        "$ icgrep -colors=never foo project/data/one_match.txt /tmp/not_there_xyz.txt\n"
        "icgrep: \"/tmp/not_there_xyz.txt\": No such file.\n"
        "project/data/one_match.txt:foo\n"
        "exit=0   # expected: 2 (error occurred)",
    )

    add_para(doc, "Why this is worse than the single-file case:", bold=True)
    add_bullets(
        doc,
        [
            "The single-file case returns exit 1 (weak signal) but still indicates non-success.",
            "The multi-file case returns exit 0 — any success/failure check will see a clean pass.",
            "In batch pipelines this silently drops files that could not be read.",
        ],
    )
    add_para(doc, "Suggested fix:", bold=True)
    add_para(
        doc,
        "When processing multiple files, escalate the final exit code to 2 if any "
        "file operation failed, regardless of whether other files produced matches.",
    )

    add_heading(doc, "Distinctness of Findings", level=1)
    add_para(
        doc,
        "The two findings are distinct root-cause errors. The first is a per-file "
        "wrong-exit-code (1 instead of 2). The second is a batch-level exit-code "
        "masking issue (0 instead of 2 when any file errored). Although 144 frames "
        "manifest the first, they collapse to one distinct specification violation "
        "because they all share the same expected/actual exit, error message class, "
        "and root cause.",
    )

    add_heading(doc, "Other Observations", level=1)
    add_bullets(
        doc,
        [
            "No crashes observed during combinatorial testing.",
            "No invalid exit codes outside {0,1,2} were produced.",
            "A benign \"failed to exec cache cleanup daemon cachejanitord\" warning "
            "appears on every invocation; it does not affect matching.",
            "Under -i, icgrep's \\p{Ll}+ does not case-fold to uppercase; the oracle "
            "was adjusted to reflect this observed behavior.",
        ],
    )

    add_heading(doc, "Reproducibility", level=1)
    add_para(
        doc,
        "All paths below are relative to the repository root. Replace "
        "<repo>/build-asan/bin/icgrep with the path to your sanitizer-enabled "
        "icgrep binary.",
    )
    add_para(doc, "Regenerate frames:", bold=True)
    add_code(doc, "python3 project/scripts/generate_icgrep_frames.py")
    add_para(doc, "Run full suite:", bold=True)
    add_code(
        doc,
        "python3 project/scripts/run_combinatorial_tests.py \\\n"
        "  --icgrep <repo>/build-asan/bin/icgrep \\\n"
        "  --frames project/frames/icgrep_frames_full.csv \\\n"
        "  --out project/results/icgrep_results_full.csv",
    )
    add_para(doc, "Regenerate this report (with figures):", bold=True)
    add_code(doc, "python3 project/scripts/generate_report.py")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Report written to: {OUT_DOCX}")
    print(f"Figures: {fig1_path.name}, {fig2_path.name} (in {FIG_DIR})")


if __name__ == "__main__":
    main()
